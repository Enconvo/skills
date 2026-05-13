from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import sys
import json
import subprocess
import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from config import get_output_dir

# Fetch live COT data
print("Running cot_tracker.py --json...")
result = subprocess.run(
    ["python3", "cot_tracker.py", "--json"], 
    capture_output=True, text=True, cwd=os.path.dirname(os.path.abspath(__file__))
)
if result.returncode != 0:
    print(f"Error running cot_tracker.py: {result.stderr}")
    exit(1)

cot_data = json.loads(result.stdout)
report_date = list(cot_data.values())[0]['date']  # Use date from first asset

# Convert date string to nice format
dt = datetime.datetime.strptime(report_date, "%Y-%m-%d")
report_date_str = dt.strftime("%B %d, %Y")

doc = Document()

# STYLE-02 Colors
PRIMARY = RGBColor(0xC8, 0x10, 0x2E) # Crimson
TEXT_DARK = RGBColor(0x33, 0x33, 0x33)
TEXT_MUTED = RGBColor(0x8C, 0x82, 0x79)
BORDER = RGBColor(0xB8, 0xB0, 0xA8)
BG_ALT = "FAF7F2" # Hex for shading
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Page Setup
for section in doc.sections:
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

# Styles
style_normal = doc.styles['Normal']
style_normal.font.name = 'Georgia'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = TEXT_DARK
style_normal.paragraph_format.line_spacing = 1.3
style_normal.paragraph_format.space_after = Pt(12)

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Rockwell'
style_h1.font.size = Pt(26)
style_h1.font.bold = True
style_h1.font.color.rgb = PRIMARY
style_h1.paragraph_format.space_before = Pt(24)
style_h1.paragraph_format.space_after = Pt(12)

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Georgia'
style_h2.font.size = Pt(18)
style_h2.font.bold = True
style_h2.font.color.rgb = TEXT_DARK
style_h2.paragraph_format.space_before = Pt(18)
style_h2.paragraph_format.space_after = Pt(6)

def add_accent_rule(doc):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C8102E')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

# Cover
doc.add_paragraph('\n\n\n')
title = doc.add_paragraph("Smart Money Positioning", style='Heading 1')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph("CFTC Commitment of Traders (COT) Executive Summary", style='Heading 2')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb = TEXT_MUTED
doc.add_paragraph(f"As of: {report_date_str}").alignment = WD_ALIGN_PARAGRAPH.CENTER

# Exec Summary
doc.add_paragraph("Executive Summary", style='Heading 1')
add_accent_rule(doc)
doc.add_paragraph(f"This report outlines the latest positioning shifts from the CFTC Commitment of Traders (COT) report released on {datetime.datetime.now().strftime('%B %d')}. We track 'Smart Money' flows across key financial and commodity futures to gauge institutional sentiment.")

doc.add_paragraph("Key Takeaways:", style='Heading 2')

# Auto-generate highlights based on biggest WoW moves
sorted_assets = sorted(((k, v) for k, v in cot_data.items() if k != '_meta'), key=lambda item: abs(item[1]['hedge_fund_wow']), reverse=True)
top_3 = sorted_assets[:3]

for asset, data in top_3:
    net = data['hedge_fund_net']
    wow = data['hedge_fund_wow']
    action = "Bought" if wow > 0 else "Sold"
    trend = "Bullish" if net > 0 else "Bearish"
    doc.add_paragraph(f"• {asset}: Hedge funds {action.lower()} {abs(wow):,} contracts (Net: {net:,}, Trend: {trend}).", style='List Bullet')

# Data Table
doc.add_paragraph("\nAsset Class Breakdown", style='Heading 1')
add_accent_rule(doc)

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Shading' # fallback

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Asset'
hdr_cells[1].text = 'Hedge Fund Net'
hdr_cells[2].text = 'WoW Change'
hdr_cells[3].text = 'Asset Mgr Net'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
    # LXML shading for header
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'C8102E')
    tcPr.append(shading)

# Sort logic: custom order
order = ["Bitcoin", "S&P 500", "Nasdaq 100", "Euro FX", "Gold", "Crude Oil"]
for asset in order:
    if asset in cot_data:
        data = cot_data[asset]
        row = table.add_row().cells
        row[0].text = asset
        row[1].text = f"{data['hedge_fund_net']:,}"
        
        wow = data['hedge_fund_wow']
        sign = "+" if wow > 0 else ""
        action = "(Buy)" if wow > 0 else "(Sell)"
        row[2].text = f"{sign}{wow:,} {action}"
        
        am_net = data.get('asset_mgr_net', 'N/A')
        if isinstance(am_net, int):
            am_net = f"{am_net:,}"
        row[3].text = str(am_net)

# Add alternate row shading
for i, row in enumerate(table.rows):
    if i == 0: continue
    if i % 2 == 0:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), BG_ALT)
            tcPr.append(shading)

output_path = str(get_output_dir() / f'COT_Smart_Money_Report_{report_date}.docx')
doc.save(output_path)
print(f"Saved to {output_path}")
print(f"REPORT_PATH:{output_path}")
